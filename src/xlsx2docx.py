from docx import Document 
import sys, getopt
import pandas as pd
import re

class xlsx_reader:

    # Parameters
    USED_COLS = ['Type','Field','Bit_Range', 'Reset_Value','Access','Description']

    def __init__(self,xlsx_file,xlsx_sheet) :
        self.wb = pd.read_excel(io=xlsx_file,sheet_name=xlsx_sheet,usecols=self.USED_COLS)
        self.__get_index_of_registers()

    def lint_check(self) :
        #TODO: add xlsx_input validation here
        pass

    def get_design_params(self) :
        param_idx  = self.wb[self.wb['Type'] == 'parameter'].index.values
        self.dsgn_name  = self.wb.iloc[param_idx[0]]['Field']
        self.addr_width = self.wb.iloc[param_idx[0]]['Bit_Range']
        self.data_width = self.wb.iloc[param_idx[0]]['Reset_Value']
        self.strb_width = int(self.data_width/8)

    def get_register_info(self,start_row,end_row) :
        reg_name    = self.wb.iloc[start_row]['Field']
        reg_fields  = self.wb.iloc[start_row +1 : end_row] 
        reg_address = self.wb.iloc[start_row]['Bit_Range']
        reg_desc    = self.wb.iloc[start_row]['Description']

        reg_reset_value = 0x0
        reg_bit_width   = 0x0
        reg_field_lst   = list()
        for i in range(0,len(reg_fields)) :
            field_name      = reg_fields.iloc[i]['Field']         
            field_bit_range = reg_fields.iloc[i]['Bit_Range']         
            field_acc_type  = reg_fields.iloc[i]['Access']        
            field_rst_value = reg_fields.iloc[i]['Reset_Value']       
            field_desc      = reg_fields.iloc[i]['Description']         

            # Extract register fields 
            field_bit_indexs = self.__get_bit_indexs(field_bit_range)
            field_bit_width  = (field_bit_indexs['msb'] - field_bit_indexs['lsb']) + 1
            # Calculate Reset_Value & Bit_Width
            int_field_rst_value = int(field_rst_value,0)
            bit_mask = (2**self.data_width-1) >> (self.data_width - field_bit_width)
            reg_reset_value += (int_field_rst_value & bit_mask) << field_bit_indexs['lsb']
            reg_bit_width += field_bit_width

            reg_field_lst.append({
                 'name'      : field_name
                ,'bit_range' : field_bit_range
                ,'access'    : field_acc_type
                ,'reset'     : field_rst_value
                ,'field_desc': field_desc
            })


        reg_info = {
             'name'        : reg_name
            ,'address'     : reg_address
            ,'bit_width'   : reg_bit_width
            ,'reset_value' : reg_reset_value
            ,'reg_desc'    : reg_desc
            ,'fields'      : reg_field_lst
        }

        return reg_info
    #---------------------------------------------
    # Private functions
    #---------------------------------------------
    def __get_bit_indexs(self,bit_range) :
        if(pd.isna(bit_range) == True) :
            lsb = 'inf'
            msb = 'inf'
        else :
            bit_indexs = list(map(int,re.findall(r'\d+',bit_range)))
            num_indexs = len(bit_indexs)
            if(num_indexs == 1): 
                bit_indexs.append(bit_indexs[0]) # append dummy value to avoid syntax error
            lsb = min(bit_indexs)
            msb = max(bit_indexs)
        index_out = {'lsb' : lsb, 'msb': msb}
        return index_out

    def __get_index_of_registers(self) :
        self.reg_start_rows = self.wb[self.wb['Type'] == 'register'].index.values
        self.reg_end_rows   = self.wb[self.wb['Type'] == 'comment'].index.values
        assert(len(self.reg_start_rows) == len(self.reg_end_rows)) , "Missing 'register' or 'comment' field !!!!"

#----------------------------------------
# Main Scripts
#----------------------------------------
def print_help() :
    print ("xlsx2docx.py -i <excel-based register input file -w <docx templete file> -o <output dir> -l <heading level>")
    
def main(argv) :

    # Terminal Arguments
    XLSX_IN   = ""
    DOCX_TEMP  = ""
    OUT_DIR   = "./"
    try :
        opts,args = getopt.getopt(argv,"hi:w:o:l:",["ifile=","wfile=","odir=","heading_level="])
    except getopt.GetoptError:
        print_help()
        sys.exit(2)
    for opt, arg in opts:
        if(opt == '-h') :
            print_help()
            sys.exit()
        elif opt in ("-i","--ifile"):
            XLSX_IN = arg
        elif opt in ("-w","--wfile"):
            DOCX_TEMP = arg
        elif opt in ("-o","--odir"):
            OUT_DIR = arg
        elif opt in ("-l","--heading_level"):
            HEADING_LEVEL = int(arg)

    csr = xlsx_reader(XLSX_IN,"register_set")

    # perform lint check for excel file
    csr.lint_check()

    # get design paramter
    csr.get_design_params()
    print("Generating CSR document for:")
    print("+ dsgn_name : " + csr.dsgn_name)
    print("+ addr_width: " + str(csr.addr_width))
    print("+ data_width: " + str(csr.data_width))
    print("+ total registers : " + str(len(csr.reg_start_rows)))

    DOCX_OUT = OUT_DIR + csr.dsgn_name.lower() + "_csr_ug.docx"
    print("Generating " + DOCX_OUT + " ...")
    document = Document(DOCX_TEMP)

    for reg_no in range(0,len(csr.reg_start_rows)) :
        reg = csr.get_register_info(csr.reg_start_rows[reg_no],csr.reg_end_rows[reg_no])
        reg_name = reg['name']
        reg_addr = reg['address']
        reg_desc = reg['reg_desc']
        reg_bw   = reg['bit_width']
        reg_rst  = reg['reset_value']
        lst_reg_fields = reg['fields']
        print("Generating " + reg_name + " ...")

        document.add_heading("Address \"" + reg_addr + "\"; Register \"" + reg_name + "\"",level=HEADING_LEVEL)
        document.add_paragraph ("Register: " + reg_desc         , style='Normal')
        document.add_paragraph ("Address : " + reg_addr         , style='Normal')
        document.add_paragraph ("Reset Value : " + hex(reg_rst) , style='Normal')

        table = document.add_table(rows=len(lst_reg_fields)+1, cols=5)
        table.style = 'Basic_table_style'
        table.cell(0, 0).text = "Field Name"
        table.cell(0, 1).text = "Bits"
        table.cell(0, 2).text = "Access Type"
        table.cell(0, 3).text = "Reset Value"
        table.cell(0, 4).text = "Description"
        for row in range(1,len(lst_reg_fields)+1):
            table.cell(row, 0).text = lst_reg_fields[row-1]['name'      ]
            table.cell(row, 1).text = lst_reg_fields[row-1]['bit_range' ]
            table.cell(row, 2).text = lst_reg_fields[row-1]['access'    ]
            table.cell(row, 3).text = lst_reg_fields[row-1]['reset'     ]
            table.cell(row, 4).text = lst_reg_fields[row-1]['field_desc']
    document.save(DOCX_OUT) # Save document

if __name__ == "__main__" :
    main(sys.argv[1:])
